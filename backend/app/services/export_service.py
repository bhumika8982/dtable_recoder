"""Export Minutes of Meeting (MOM) to PDF (reportlab) and DOCX (python-docx).

Both builders return raw bytes so the router can stream them directly.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


def build_pdf(title: str, mom: dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=title)
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(title, styles["Title"]), Spacer(1, 0.2 * inch)]

    if mom.get("summary"):
        story.append(Paragraph("Overview", styles["Heading2"]))
        story.append(Paragraph(mom["summary"], styles["BodyText"]))
        story.append(Spacer(1, 0.15 * inch))

    def bullet_section(heading: str, lines: list[str]) -> None:
        if not lines:
            return
        story.append(Paragraph(heading, styles["Heading2"]))
        story.append(
            ListFlowable(
                [ListItem(Paragraph(str(x), styles["BodyText"])) for x in lines],
                bulletType="bullet",
            )
        )
        story.append(Spacer(1, 0.15 * inch))

    bullet_section("Attendees", mom.get("attendees", []))
    bullet_section("Key Discussion Points", mom.get("key_points", []))
    bullet_section("Action Items", mom.get("action_items", []))
    bullet_section("Next Steps", mom.get("next_steps", []))

    doc.build(story)
    return buf.getvalue()


def build_docx(title: str, mom: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading(title, level=0)

    if mom.get("summary"):
        document.add_heading("Overview", level=1)
        document.add_paragraph(mom["summary"])

    def bullet_section(heading: str, lines: list[str]) -> None:
        if not lines:
            return
        document.add_heading(heading, level=1)
        for line in lines:
            document.add_paragraph(str(line), style="List Bullet")

    bullet_section("Attendees", mom.get("attendees", []))
    bullet_section("Key Discussion Points", mom.get("key_points", []))
    bullet_section("Action Items", mom.get("action_items", []))
    bullet_section("Next Steps", mom.get("next_steps", []))

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
